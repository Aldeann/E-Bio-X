"""Extract plain text from uploaded material files (PDF, DOCX, TXT).

Used by the AI explanation knowledge base to turn uploaded files into
retrievable context. Extraction is best-effort: any failure returns ''.
"""
import io
import os
import re


def _clean_text(text):
    if not text:
        return ''
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def extract_pdf_text(path_or_bytes):
    """path: file path on disk. Returns extracted text or ''."""
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(path_or_bytes)
        return _clean_text(text)
    except Exception:
        return ''


def extract_docx_text(path_or_bytes):
    """path: file path on disk. Returns extracted text (paragraphs+tables) or ''."""
    try:
        from docx import Document
        doc = Document(path_or_bytes)
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return _clean_text(' '.join(parts))
    except Exception:
        return ''


def extract_file_text(path, file_type=None):
    """Best-effort text extraction from a uploaded file on disk.

    file_type is a lowercase extension ('pdf', 'docx', 'doc', 'txt').
    """
    if not path or not os.path.isfile(path):
        return ''
    ext = (file_type or os.path.splitext(path)[1].lstrip('.')).lower()
    try:
        if ext == 'pdf':
            return extract_pdf_text(path)
        if ext in ('docx', 'doc'):
            return extract_docx_text(path)
        if ext in ('txt', 'md'):
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                return _clean_text(f.read())
    except Exception:
        return ''
    return ''


def extract_upload_bytes(data, filename):
    """Extract text from uploaded in-memory bytes (werkzeug FileStorage)."""
    ext = (filename.rsplit('.', 1)[-1] if '.' in filename else '').lower()
    try:
        if ext == 'pdf':
            return extract_pdf_text(io.BytesIO(data))
        if ext in ('docx', 'doc'):
            return extract_docx_text(io.BytesIO(data))
        if ext in ('txt', 'md'):
            return _clean_text(data.decode('utf-8', errors='ignore'))
    except Exception:
        return ''
    return ''